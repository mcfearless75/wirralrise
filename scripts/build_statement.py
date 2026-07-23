"""
Wirral Rise — Statement of Services generator.
Produces: wirralrise-statement-of-services.docx in the repo root.
Run: python scripts/build_statement.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "..", "wirralrise-statement-of-services.docx")

# ── Colours ─────────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1B, 0x3A, 0x5C)
BLUE       = RGBColor(0x29, 0x80, 0xB9)
TEAL       = RGBColor(0x17, 0xA5, 0x89)
AMBER      = RGBColor(0xE6, 0x7E, 0x22)
YELLOW     = RGBColor(0xF4, 0xD0, 0x3F)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF4, 0xF7, 0xFB)
TEXT_MID   = RGBColor(0x37, 0x41, 0x51)
TEXT_MUTED = RGBColor(0x6B, 0x72, 0x80)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 6)))
            el.set(qn("w:color"), val.get("color", "1B3A5C"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para(doc, text, size=11, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = "Calibri"
    return p

def heading(doc, text, level=1, color=None, space_before=14, space_after=4):
    sizes = {1: 22, 2: 16, 3: 13, 4: 11}
    c     = color or NAVY
    return para(doc, text, size=sizes.get(level, 11), bold=True, color=c,
                space_before=space_before, space_after=space_after)

def bullet(doc, text, indent=0.6, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent    = Cm(indent)
    p.paragraph_format.space_before   = Pt(1)
    p.paragraph_format.space_after    = Pt(3)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.name  = "Calibri"
    if color:
        run.font.color.rgb = color
    return p

def rule(doc, color="1B3A5C", size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def colour_block(doc, title, body_lines, bg_hex="EBF5FB", title_color=None, accent_color=None):
    """Single-column table used as a coloured callout block."""
    t_color = title_color or NAVY
    a_color = accent_color or TEAL
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = t_color
    r.font.name = "Calibri"
    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.left_indent  = Cm(0.3)
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after  = Pt(2)
        br = bp.add_run(f"• {line}")
        br.font.size  = Pt(11)
        br.font.name  = "Calibri"
        br.font.color.rgb = TEXT_MID
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def service_row(doc, title, description, icon_char="▶"):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(1.2)
    tbl.columns[1].width = Cm(14.8)
    ic = tbl.cell(0, 0)
    set_cell_bg(ic, "EBF5FB")
    ip = ic.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = ip.add_run(icon_char)
    ir.font.size  = Pt(16)
    ir.font.color.rgb = TEAL
    ic._tc.get_or_add_tcPr()  # ensure exists
    tc = tbl.cell(0, 1)
    set_cell_bg(tc, "FFFFFF")
    tp = tc.paragraphs[0]
    tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.space_after  = Pt(2)
    tr_title = tp.add_run(f"{title}\n")
    tr_title.font.bold  = True
    tr_title.font.size  = Pt(12)
    tr_title.font.color.rgb = NAVY
    tr_title.font.name  = "Calibri"
    tr_desc = tp.add_run(description)
    tr_desc.font.size   = Pt(11)
    tr_desc.font.color.rgb = TEXT_MID
    tr_desc.font.name   = "Calibri"
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─────────────────────────────────────────────────────────────────
# COVER BLOCK  (coloured table)
# ─────────────────────────────────────────────────────────────────
cover = doc.add_table(rows=1, cols=1)
cover.style = "Table Grid"
cc = cover.cell(0, 0)
set_cell_bg(cc, "1B3A5C")

title_p = cc.paragraphs[0]
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(28)
title_p.paragraph_format.space_after  = Pt(4)
tr = title_p.add_run("Wirral Rise")
tr.font.size  = Pt(36)
tr.font.bold  = True
tr.font.color.rgb = WHITE
tr.font.name  = "Calibri"

sub_p = cc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after  = Pt(4)
sr = sub_p.add_run("Statement of Services")
sr.font.size  = Pt(20)
sr.font.bold  = False
sr.font.color.rgb = YELLOW
sr.font.name  = "Calibri"

tag_p = cc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_p.paragraph_format.space_before = Pt(0)
tag_p.paragraph_format.space_after  = Pt(28)
tgr = tag_p.add_run("Specialist Alternative Provision for Neurodivergent Children | Wirral")
tgr.font.size  = Pt(12)
tgr.font.color.rgb = RGBColor(0xB2, 0xD3, 0xEE)
tgr.font.name  = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# ─────────────────────────────────────────────────────────────────
# CONTACT STRIP
# ─────────────────────────────────────────────────────────────────
ct = doc.add_table(rows=1, cols=3)
ct.style = "Table Grid"
ct.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, val) in enumerate([
    ("📍  Address", "143–145 Ford Road, Upton Village, Wirral, CH49 0XB"),
    ("📞  Phone",   "+44 (0)7515 117816"),
    ("✉  Email",   "admin@wirralrise.com"),
]):
    c = ct.cell(0, i)
    set_cell_bg(c, "EBF5FB")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    lr = p.add_run(f"{label}\n")
    lr.font.bold  = True
    lr.font.size  = Pt(9)
    lr.font.color.rgb = NAVY
    lr.font.name  = "Calibri"
    vr = p.add_run(val)
    vr.font.size  = Pt(10)
    vr.font.color.rgb = TEXT_MID
    vr.font.name  = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(6)
rule(doc)

# ─────────────────────────────────────────────────────────────────
# 1. ABOUT US
# ─────────────────────────────────────────────────────────────────
heading(doc, "1.  About Wirral Rise", level=1, space_before=12)
para(doc,
    "Wirral Rise is a specialist alternative provision based in Upton Village, Wirral, "
    "established to support neurodivergent children in Years 1–6 who are unable to access "
    "mainstream education due to anxiety, Emotionally Based School Avoidance (EBSA), SEND or "
    "trauma-related difficulties. We provide a warm, structured and therapeutic environment "
    "where children feel safe, valued and genuinely understood.",
    color=TEXT_MID, space_after=6)
para(doc,
    "Our fundamental aim is always to help each child develop the confidence, regulation and "
    "resilience they need to return to mainstream school when the time is right — at their pace, "
    "not ours.",
    color=TEXT_MID, space_after=8)

# ─────────────────────────────────────────────────────────────────
# 2. OUR APPROACH
# ─────────────────────────────────────────────────────────────────
heading(doc, "2.  Our Approach", level=1, space_before=10)
para(doc,
    "Wirral Rise operates a Trauma-Informed, Child-Led model of provision. Everything we do "
    "is underpinned by three core principles:",
    color=TEXT_MID, space_after=6)

colour_block(doc,
    "Core Principles",
    [
        "Child-Led, Interest-Based Learning — children help shape what they explore, which "
          "increases engagement, confidence and retention.",
        "Therapeutic Relationships First — trust must be established before learning can happen. "
          "We invest time in genuine relationships.",
        "Planned Pathway Back to School — provision is always designed with the child's return "
          "to mainstream as the goal.",
    ],
    bg_hex="E8F8F5", title_color=TEAL)

# ─────────────────────────────────────────────────────────────────
# 3. SERVICES
# ─────────────────────────────────────────────────────────────────
heading(doc, "3.  Services Provided", level=1, space_before=10)
para(doc,
    "The following specialist services are embedded into daily provision at Wirral Rise. "
    "They are not separate, clinical add-ons — they are woven into the fabric of each child's "
    "everyday experience.",
    color=TEXT_MID, space_after=8)

services = [
    ("EBSA Support",
     "We specialise in Emotionally Based School Avoidance. Our practitioners understand the "
     "anxiety behind non-attendance and work patiently and compassionately to help children "
     "feel safe enough to engage with learning — and eventually return to school."),
    ("SEND & EHCP Provision",
     "Every child's individual programme is built directly from their Education, Health and "
     "Care Plan (Sections E and F). We work in close partnership with schools, local "
     "authorities and families to ensure all statutory needs are fully and properly met."),
    ("Neurodevelopmental Programmes",
     "Our specialist neurodevelopmental practitioners deliver tailored programmes for children "
     "with autism, ADHD, PDA, demand avoidance and sensory processing differences, aligned to "
     "best-practice frameworks."),
    ("Speech & Language Therapy",
     "Our specialist Speech and Language Therapist is embedded into daily provision. Therapy "
     "is woven into every session — not a separate appointment — to ensure natural, meaningful "
     "impact throughout the child's day."),
    ("Emotional Regulation & SEMH",
     "We equip children with practical tools to understand and manage their own emotions. "
     "A calmer, more regulated child is a child who is ready to learn — this sits at the "
     "heart of everything we do at Wirral Rise."),
    ("Academic Learning",
     "Qualified teachers deliver personalised literacy, numeracy and broad curriculum learning, "
     "shaped around each child's interests, strengths and pace. Learning plans are aligned to "
     "national curriculum expectations and EHCP outcomes."),
    ("Therapeutic Dog Programme",
     "Cookie, our trained therapy dog, is a core member of the Wirral Rise team. Her calming "
     "presence supports children's emotional regulation from the moment they arrive, and plays "
     "a central role in building a sense of safety and belonging."),
]

icons = ["🛡", "📋", "🧠", "💬", "💛", "📚", "🐾"]
for (title, desc), icon in zip(services, icons):
    service_row(doc, title, desc, icon_char=icon)

# ─────────────────────────────────────────────────────────────────
# 4. WHO WE SUPPORT
# ─────────────────────────────────────────────────────────────────
heading(doc, "4.  Who We Support", level=1, space_before=10)
para(doc,
    "Wirral Rise accepts referrals for children who meet one or more of the following criteria:",
    color=TEXT_MID, space_after=6)

colour_block(doc,
    "Referral Criteria",
    [
        "Primary-age children in Years 1–6 (ages 5–11)",
        "Neurodivergent profile: autism, ADHD, PDA, sensory processing differences",
        "Emotionally Based School Avoidance (EBSA) or significant anxiety preventing attendance",
        "Trauma-related difficulties affecting access to education",
        "SEMH needs requiring a specialist environment",
        "Children with an active or pending EHCP",
        "Children requiring short-term, intensive support prior to a school transition",
    ],
    bg_hex="FFF9E6", title_color=RGBColor(0xB7, 0x77, 0x0D))

para(doc,
    "We currently provide specialist provision for children across Wirral and the wider "
    "Merseyside area. LA-funded placements are considered, and we are happy to support "
    "families in navigating the funding process.",
    color=TEXT_MID, space_after=8)

# ─────────────────────────────────────────────────────────────────
# 5. THE TEAM
# ─────────────────────────────────────────────────────────────────
heading(doc, "5.  Our Team", level=1, space_before=10)
para(doc,
    "Wirral Rise is staffed by a qualified, experienced multidisciplinary team (MDT):",
    color=TEXT_MID, space_after=6)

team_tbl = doc.add_table(rows=6, cols=3)
team_tbl.style = "Table Grid"
team_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ["Name", "Role", "Qualification / Specialism"]
for j, h in enumerate(headers):
    c = team_tbl.cell(0, j)
    set_cell_bg(c, "1B3A5C")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(h)
    r.font.bold  = True
    r.font.size  = Pt(11)
    r.font.color.rgb = WHITE
    r.font.name  = "Calibri"

members = [
    ("Jayne Foley",  "Director / Lead Practitioner",         "BSc Hons, Neurodevelopmental Specialist"),
    ("Mike Jones",   "Education Lead",                       "Qualified Teacher Status (QTS)"),
    ("Sam McHugh",   "Speech & Language Therapist",          "BSc Hons Speech & Language Therapy, HCPC Registered"),
    ("Cath Downs",   "SEMH Practitioner",                    "BA Hons, SEMH Specialist"),
    ("Cookie",       "Therapy Dog",                          "Fully trained; Wirral Rise's secret weapon"),
]

for i, (name, role, qual) in enumerate(members):
    bg = "F4F7FB" if i % 2 == 0 else "FFFFFF"
    row_data = [name, role, qual]
    for j, val in enumerate(row_data):
        c = team_tbl.cell(i + 1, j)
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(val)
        r.font.size  = Pt(11)
        r.font.name  = "Calibri"
        r.font.color.rgb = TEXT_MID
        if j == 0:
            r.font.bold = True
            r.font.color.rgb = NAVY

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─────────────────────────────────────────────────────────────────
# 6. REFERRAL PROCESS
# ─────────────────────────────────────────────────────────────────
heading(doc, "6.  Referral Process", level=1, space_before=10)

steps = [
    ("Initial Enquiry",
     "Contact us by phone or email. We'll have an informal, no-pressure conversation to "
     "understand your child's needs and see if Wirral Rise is the right fit."),
    ("Assessment Visit",
     "We invite the child and family for a relaxed visit to our hub. This helps us understand "
     "the child's needs, interests and current presentation, and lets the family see the "
     "provision in action."),
    ("Individual Programme Design",
     "We develop a detailed, personalised learning and support plan aligned to the child's "
     "EHCP (Sections E and F) and agreed outcomes. This is produced in partnership with the "
     "family, school and LA."),
    ("Placement Commencement",
     "The child begins provision at an agreed pace — we always follow the child's lead on "
     "transition speed. Regular reviews are scheduled from the outset."),
    ("Ongoing Review & Reporting",
     "We provide written progress reports, contribute to Annual Review meetings and maintain "
     "open, regular communication with families, schools and the Local Authority throughout "
     "the placement."),
]

for i, (title, desc) in enumerate(steps, 1):
    srow = doc.add_table(rows=1, cols=2)
    srow.style = "Table Grid"
    srow.columns[0].width = Cm(1.0)
    srow.columns[1].width = Cm(15.0)
    nc = srow.cell(0, 0)
    set_cell_bg(nc, "17A589")
    np_ = nc.paragraphs[0]
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr_ = np_.add_run(str(i))
    nr_.font.bold  = True
    nr_.font.size  = Pt(14)
    nr_.font.color.rgb = WHITE
    nr_.font.name  = "Calibri"
    tc2 = srow.cell(0, 1)
    set_cell_bg(tc2, "E8F8F5")
    tp2 = tc2.paragraphs[0]
    tp2.paragraph_format.space_before = Pt(3)
    tp2.paragraph_format.space_after  = Pt(3)
    ttr = tp2.add_run(f"{title}\n")
    ttr.font.bold  = True
    ttr.font.size  = Pt(12)
    ttr.font.color.rgb = NAVY
    ttr.font.name  = "Calibri"
    tdr = tp2.add_run(desc)
    tdr.font.size  = Pt(11)
    tdr.font.color.rgb = TEXT_MID
    tdr.font.name  = "Calibri"
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ─────────────────────────────────────────────────────────────────
# 7. FUNDING
# ─────────────────────────────────────────────────────────────────
heading(doc, "7.  Funding & Placements", level=1, space_before=10)
para(doc,
    "Wirral Rise accepts placements funded by:",
    color=TEXT_MID, space_after=6)
for line in [
    "Local Authority EHCP funding (Wirral LA and neighbouring authorities)",
    "School-funded packages (agreed directly with the placing school)",
    "Private / family-funded placements",
]:
    bullet(doc, line, color=TEXT_MID)

para(doc,
    "\nWe work openly and collaboratively with placing authorities and are experienced in "
    "contributing to Annual Reviews, EHCP amendments and Local Offer requirements. "
    "Please contact us to discuss individual funding arrangements.",
    color=TEXT_MID, space_after=8)

# ─────────────────────────────────────────────────────────────────
# 8. SAFEGUARDING
# ─────────────────────────────────────────────────────────────────
heading(doc, "8.  Safeguarding & Policies", level=1, space_before=10)
colour_block(doc,
    "Safeguarding Commitment",
    [
        "Wirral Rise is committed to safeguarding and promoting the welfare of children.",
        "All staff hold current Enhanced DBS certificates.",
        "A Designated Safeguarding Lead (DSL) is on site at all times.",
        "All policies — including Safeguarding, SEND, Behaviour and Health & Safety — "
          "are reviewed annually and are available on request.",
        "Wirral Rise operates in accordance with Keeping Children Safe in Education (KCSiE) "
          "and Working Together to Safeguard Children.",
    ],
    bg_hex="E9F7EF", title_color=RGBColor(0x1E, 0x8B, 0x49))

# ─────────────────────────────────────────────────────────────────
# 9. CONTACT
# ─────────────────────────────────────────────────────────────────
rule(doc)
heading(doc, "Contact Us", level=2, space_before=8)
para(doc,
    "We'd love to hear from you. Whether you're a parent, a school or a local authority "
    "professional — please get in touch for a friendly, no-obligation conversation.",
    color=TEXT_MID, space_after=6)

for label, val in [
    ("Address",  "143–145 Ford Road, Upton Village, Wirral, CH49 0XB"),
    ("Phone",    "+44 (0)7515 117816"),
    ("Email",    "admin@wirralrise.com"),
    ("Website",  "www.wirralrise.co.uk"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    lr = p.add_run(f"{label}: ")
    lr.font.bold  = True
    lr.font.size  = Pt(11)
    lr.font.color.rgb = NAVY
    lr.font.name  = "Calibri"
    vr = p.add_run(val)
    vr.font.size  = Pt(11)
    vr.font.color.rgb = TEXT_MID
    vr.font.name  = "Calibri"

# footer note
doc.add_paragraph()
rule(doc, color="CCCCCC", size=6)
fn = para(doc,
    "© 2024 Wirral Rise. This document is provided for information purposes. "
    "All policies and procedures are available on request.",
    size=9, color=TEXT_MUTED, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)

# ─────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"✅  Saved: {OUTPUT}")
